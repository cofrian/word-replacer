import streamlit as st
import os
import zipfile
import shutil
from docx import Document

def reemplazar_en_documento(ruta_entrada, ruta_salida, reemplazos):
    doc = Document(ruta_entrada)

    # Reemplazar en párrafos
    for p in doc.paragraphs:
        for buscar, reemplazar in reemplazos.items():
            if buscar in p.text:
                p.text = p.text.replace(buscar, reemplazar)

    # Reemplazar en tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for buscar, reemplazar in reemplazos.items():
                    if buscar in celda.text:
                        celda.text = celda.text.replace(buscar, reemplazar)

    # Guardar documento modificado
    doc.save(ruta_salida)

# --- Streamlit UI ---
st.title("🔄 Reemplazo Masivo en Word (.docx)")

archivo_zip = st.file_uploader("📦 Sube un archivo ZIP con documentos Word (.docx)", type="zip")

st.markdown("✏️ **Agrega pares de texto a buscar y reemplazar**")
reemplazos = {}
num_pares = st.number_input("Número de pares búsqueda/reemplazo", min_value=1, max_value=20, value=1, step=1)

for i in range(num_pares):
    buscar = st.text_input(f"🔎 Buscar texto #{i+1}", key=f"buscar_{i}")
    reemplazar = st.text_input(f"✏️ Reemplazar por #{i+1}", key=f"reemplazar_{i}")
    if buscar and reemplazar:
        reemplazos[buscar] = reemplazar

if st.button("🚀 Procesar documentos"):
    if archivo_zip is None:
        st.error("❌ Debes subir un archivo ZIP con documentos .docx")
    elif not reemplazos:
        st.error("❌ Debes añadir al menos un par búsqueda/reemplazo")
    else:
        with st.spinner("⏳ Procesando documentos..."):
            # Crear carpetas temporales
            temp_input = "temp_input"
            temp_output = "temp_output"
            os.makedirs(temp_input, exist_ok=True)
            os.makedirs(temp_output, exist_ok=True)

            # Extraer ZIP subido
            with zipfile.ZipFile(archivo_zip, 'r') as zip_ref:
                zip_ref.extractall(temp_input)

            # Procesar documentos en todas las carpetas y subcarpetas
            for root, dirs, files in os.walk(temp_input):
                for archivo in files:
                    if archivo.endswith(".docx"):
                        ruta_docx = os.path.join(root, archivo)

                        # Mantener estructura de carpetas en salida
                        relative_path = os.path.relpath(root, temp_input)
                        output_dir = os.path.join(temp_output, relative_path)
                        os.makedirs(output_dir, exist_ok=True)

                        nombre_modificado = f"MOD_{archivo}"
                        ruta_modificado = os.path.join(output_dir, nombre_modificado)

                        reemplazar_en_documento(ruta_docx, ruta_modificado, reemplazos)

            # Crear ZIP con resultados y mantener estructura de carpetas
            resultado_zip = "resultado.zip"
            with zipfile.ZipFile(resultado_zip, 'w') as zipf:
                for root, dirs, files in os.walk(temp_output):
                    for file in files:
                        abs_file_path = os.path.join(root, file)
                        relative_file_path = os.path.relpath(abs_file_path, temp_output)
                        zipf.write(abs_file_path, relative_file_path)

            # Mostrar enlace de descarga
            with open(resultado_zip, "rb") as f:
                st.download_button("⬇️ Descargar Word modificados (ZIP)", f, file_name="resultado.zip")

            # Limpiar carpetas temporales
            shutil.rmtree(temp_input)
            shutil.rmtree(temp_output)
            os.remove(resultado_zip)

            st.success("🎉 Procesamiento completo: Word modificados listos para descargar")
